from .ollama_client import query_ollama, EXTRACTION_PROMPT_TEMPLATE, CLASSIFY_PROMPT_TEMPLATE, ROLE_PROMPT_TEMPLATE, OUR_COMPANY
import json
import logging
import re
from io import BytesIO
from typing import Optional

MAX_CHARS = 2000  # Максимальная длина текста для LLM (уменьшено для ускорения)
OVERLAP = 500     # Перекрытие между окнами (уменьшено)

# --- Классификация типа документа через LLM ---
# CLASSIFY_PROMPT_TEMPLATE теперь импортируется из ollama_client.py

def classify_document_llm(text: str) -> str:
    prompt = CLASSIFY_PROMPT_TEMPLATE.format(text=text[:2000])
    logging.info(f"Prompt to LLM for classification: {prompt}")
    try:
        response = query_ollama(prompt)
        # Берём только первое слово из ответа
        doc_type = response.strip().split()[0].lower()
        return doc_type
    except Exception as e:
        logging.error(f"Error classifying document: {e}")
        return "иной"

# --- Извлечение текста для разных типов документов ---
def extract_full_text_from_pdf(file_path):
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            logging.info(f"[PDF] Извлечённый текст (первые 200 символов): {text[:200]}")
            if text.strip():
                return text
            # Если текст пустой — пробуем OCR
            logging.info("[PDF] Текст не найден, пробую OCR для сканированного PDF...")
            from PIL import Image
            import pytesseract
            import cv2
            import numpy as np
            import traceback
            import subprocess
            import os
            ocr_text = ""
            for i, page in enumerate(pdf.pages):
                img = page.to_image(resolution=400).original
                buf = BytesIO()
                img.save(buf, format='PNG')
                buf.seek(0)
                pil_img = Image.open(buf)
                img_array = np.array(pil_img)
                if len(img_array.shape) == 3:
                    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                else:
                    gray = img_array
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
                gray = clahe.apply(gray)
                gray = cv2.medianBlur(gray, 3)
                enhanced_img = Image.fromarray(gray)
                # Сохраняем PNG для диагностики
                png_path = f"/tmp/ocr_page_{i+1}.png"
                enhanced_img.save(png_path)
                logging.info(f"[PDF][OCR] Saved {png_path} for manual OCR test")
                custom_config = r'--oem 3 --psm 6'
                logging.info(f"[PDF][OCR] config: {custom_config}")
                try:
                    # Сначала пробуем pytesseract
                    page_text = pytesseract.image_to_string(enhanced_img, lang='rus', config=custom_config)
                except Exception as ocr_e:
                    logging.error(f"[PDF][OCR] Ошибка pytesseract: {ocr_e}")
                    logging.error(traceback.format_exc())
                    # Пробуем через subprocess
                    try:
                        result = subprocess.run([
                            'tesseract', png_path, 'stdout', '-l', 'rus', '--oem', '3', '--psm', '6'
                        ], capture_output=True, text=True)
                        logging.error(f"[PDF][OCR][subprocess] stderr: {result.stderr}")
                        page_text = result.stdout
                    except Exception as sub_e:
                        logging.error(f"[PDF][OCR][subprocess] Ошибка: {sub_e}")
                        logging.error(traceback.format_exc())
                        page_text = ''
                ocr_text += page_text + "\n"
            logging.info(f"[PDF][OCR] Извлечённый текст (первые 200 символов): {ocr_text[:200]}")
            return ocr_text.strip()
    except Exception as e:
        import traceback
        logging.error(f"[PDF] Ошибка при извлечении текста: {e}")
        logging.error(traceback.format_exc())
        return ""

def extract_text_from_pdf_contract(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            first_page_text = pdf.pages[0].extract_text() or ""
            requisites_text = ""
            for page in pdf.pages:
                text = page.extract_text() or ""
                if "реквизит" in text.lower():
                    requisites_text += text + "\n"
            return (first_page_text + "\n" + requisites_text).strip()
    except Exception:
        return ""

def extract_full_text_from_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""

def extract_text_from_docx_contract(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        first_paragraphs = [p.text for p in doc.paragraphs[:20] if p.text.strip()]
        requisites = [p.text for p in doc.paragraphs if "реквизит" in p.text.lower()]
        return ("\n".join(first_paragraphs) + "\n" + "\n".join(requisites)).strip()
    except Exception:
        return ""

def extract_text_from_jpg(file_path):
    try:
        from PIL import Image
        import pytesseract
        import cv2
        import numpy as np
        import traceback
        pil_img = Image.open(file_path)
        img_array = np.array(pil_img)
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        gray = clahe.apply(gray)
        gray = cv2.medianBlur(gray, 3)
        enhanced_img = Image.fromarray(gray)
        custom_config = r'--oem 3 --psm 6'
        logging.info(f"[JPG][OCR] config: {custom_config}")
        try:
            return pytesseract.image_to_string(enhanced_img, lang='rus', config=custom_config)
        except Exception as ocr_e:
            logging.error(f"[JPG][OCR] Ошибка pytesseract: {ocr_e}")
            logging.error(traceback.format_exc())
            return ""
    except Exception as e:
        import traceback
        logging.error(f"[JPG] Ошибка при извлечении текста: {e}")
        logging.error(traceback.format_exc())
        return ""

# --- Универсальная эвристика для классификации типа документа ---
def classify_document_universal(text: str) -> str:
    text_lower = text.lower()
    if "упд" in text_lower or "универсальный передаточный" in text_lower:
        return "упд"
    if "акт" in text_lower:
        return "акт"
    if "накладная" in text_lower:
        return "накладная"
    if "счёт-фактура" in text_lower or "счет-фактура" in text_lower:
        return "счёт-фактура"
    if "передаточный" in text_lower:
        return "передаточный"
    if "счет" in text_lower or "счёт" in text_lower:
        return "счёт"
    if "договор" in text_lower:
        return "договор"
    if "invoice" in text_lower:
        return "счёт"
    if "contract" in text_lower:
        return "договор"
    if re.search(r"\bакт[а-я]*\b", text_lower):
        return "акт"
    if re.search(r"\bдоговор[а-я]*\b", text_lower):
        return "договор"
    if re.search(r"\bнакладн[а-я]*\b", text_lower):
        return "накладная"
    if re.search(r"\bупд\b", text_lower):
        return "упд"
    return "иной"

# --- Список нужных полей для каждого типа документа ---
def get_fields_for_doc_type(doc_type: str):
    doc_type = (doc_type or "").lower()
    if doc_type == "акт":
        return ["counterparty", "date", "amount", "doc_number"]
    if doc_type == "договор":
        return ["counterparty", "date", "doc_number"]
    if doc_type == "счёт":
        return ["counterparty", "date", "amount", "doc_number"]
    if doc_type == "упд":
        return ["counterparty", "date", "amount", "doc_number", "inn"]
    if doc_type == "накладная":
        return ["counterparty", "date", "amount", "doc_number"]
    return ["counterparty", "date", "amount", "doc_number", "inn", "contract_number"]

# --- Универсальная функция для bot/main.py ---
def process_file_with_classification(file_path):
    ext = file_path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        full_text = extract_full_text_from_pdf(file_path)
        doc_type = classify_document_universal(full_text[:MAX_CHARS])
        logging.info(f"Document type (universal): {doc_type}")
        if doc_type == "договор":
            return extract_text_from_pdf_contract(file_path)
        else:
            return full_text[:MAX_CHARS]
    elif ext == "docx":
        full_text = extract_full_text_from_docx(file_path)
        doc_type = classify_document_universal(full_text[:MAX_CHARS])
        logging.info(f"Document type (universal): {doc_type}")
        if doc_type == "договор":
            return extract_text_from_docx_contract(file_path)
        else:
            return full_text[:MAX_CHARS]
    elif ext in ("jpg", "jpeg"):
        return extract_text_from_jpg(file_path)
    else:
        return None


def clean_text(text: str) -> str:
    # Удаляем лишние пробелы, пустые строки, оставляем только буквы, цифры, знаки препинания
    text = re.sub(r'[^\w\d\s.,:;!?@#№\-_/\\()\[\]{}"\'\n]', '', text, flags=re.UNICODE)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return ' '.join(lines)


def merge_fields(base: dict, new: dict) -> dict:
    # Объединяет два результата, не перезаписывая найденные значения
    result = base.copy()
    for k, v in new.items():
        if (not result.get(k)) or result[k] == "-" or result[k].lower() in ("not specified", "none", "-"):
            if v and v != "-" and v.lower() not in ("not specified", "none", "-"):
                result[k] = v
    return result


def determine_company_role(text: str) -> str:
    """
    Определяет роль нашей компании в документе.
    Возвращает: 'поставщик', 'покупатель' или 'не указана'.
    """
    try:
        prompt = ROLE_PROMPT_TEMPLATE.format(text=text[:MAX_CHARS], our_company=OUR_COMPANY)
        logging.info(f"Prompt to LLM for role determination: {prompt[:200]}...")
        response = query_ollama(prompt)
        role = response.strip().split()[0].lower()
        logging.info(f"Company role determined: {role}")
        return role
    except Exception as e:
        logging.error(f"Error determining company role: {e}")
        return "не указана"

def is_suspicious(value, field):
    if not value or value == "-":
        return True
    if field == "amount":
        try:
            normalized = re.sub(r"[^0-9.,]", "", value)
            normalized = normalized.replace(" ", "").replace("\xa0", " ")
            # Если и точка, и запятая — считаем, что запятая разделяет копейки
            if "," in normalized and "." in normalized:
                # Убираем пробелы тысяч, оставляем запятую как десятичный
                normalized = normalized.replace(".", "").replace(",", ".")
            else:
                normalized = normalized.replace(",", ".")
            amount = float(normalized)
            return amount < 10
        except Exception:
            return True
    if field == "doc_number":
        token = value.strip()
        return len(token) < 3 or not re.search(r"[0-9]", token)
    if field == "date":
        # Можно добавить проверку на слишком старую/будущую дату
        return False
    return False

# --- Быстрый путь для извлечения ключевых полей ---
def extract_fields_from_text(doc_text: str, rag_context: Optional[list] = None, doc_type: Optional[str] = None) -> dict:
    """
    Сначала пытаемся извлечь ключевые поля регулярками/паттернами (быстрый путь).
    Если не удалось — используем LLM (медленный путь).
    doc_type: если передан, используется для контекстного поиска даты и других полей
    """
    import re
    clean = clean_text(doc_text)
    total_len = len(clean)
    result = {k: "-" for k in ["inn", "counterparty", "doc_number", "date", "amount", "subject", "contract_number"]}

    # --- Контекстный поиск даты (prioritize "от" после типа/заголовка) ---
    date_patterns = [r"\b\d{2}[./]\d{2}[./]\d{4}\b", r"\b\d{2} [а-я]+ \d{4}\b"]
    date_candidates = []
    for pat in date_patterns:
        for m in re.finditer(pat, clean):
            date_candidates.append((m.start(), m.group(0)))
    # Если есть doc_type, ищем дату рядом с ключевым словом
    if doc_type and date_candidates:
        doc_type_keywords = {
            "акт": ["акт", "акта"],
            "договор": ["договор", "contract"],
            "счёт": ["счёт", "счет", "invoice"],
            "упд": ["упд", "универсальный передаточный"],
            "накладная": ["накладная"]
        }
        keywords = doc_type_keywords.get(doc_type.lower(), [])
        best = None
        best_dist = 99999
        for kw in keywords:
            for m in re.finditer(kw, clean.lower()):
                kw_pos = m.start()
                for date_pos, date_val in date_candidates:
                    dist = abs(date_pos - kw_pos)
                    if dist < best_dist:
                        best_dist = dist
                        best = date_val
        if best:
            result["date"] = best
        else:
            result["date"] = date_candidates[0][1]
    elif date_candidates:
        result["date"] = date_candidates[0][1]
    # Быстрый путь: регулярки для ИНН, даты, суммы, номера документа
    # ИНН (10 или 12 цифр)
    # Исключаем случайные 10-12-значные номера счёта/телефона, ищем рядом с ключами
    inn_match = None
    for m in re.finditer(r"\b\d{10}\b|\b\d{12}\b", clean):
        left = max(0, m.start()-20)
        ctx = clean[left:m.start()].lower()
        if any(k in ctx for k in ["инн", "inn", "налогопл"]):
            inn_match = m
            break
    if inn_match:
        result["inn"] = inn_match.group(0)
    # Сумма: выбираем кандидата из строк с якорями
    amount_candidates = []
    for m in re.finditer(r"\b\d{1,3}(?:[\s\.\,]\d{3})*(?:[\.,]\d{2})?\b", clean):
        # Окно вокруг числа
        left = max(0, m.start()-30)
        right = min(len(clean), m.end()+30)
        ctx = clean[left:right].lower()
        score = 0
        if any(k in ctx for k in ["итого", "всего к оплате", "к оплате", "сумма к оплате", "amount due"]):
            score += 3
        if any(k in ctx for k in ["с ндс", "без ндс", "nds", "налог"]):
            score += 1
        # штраф, если рядом "шт", "кол-во"
        if any(k in ctx for k in [" шт", "кол-во", "ед."]):
            score -= 1
        amount_candidates.append((score, m.group(0)))
    if amount_candidates:
        amount_candidates.sort(key=lambda x: x[0], reverse=True)
        result["amount"] = amount_candidates[0][1]
    # Номер документа: приоритет близко к символу № или словам "Счёт", "Акт", "Накладная"
    docnum_candidates = []
    # Варианты: "№ 123", "N123", после ключевых слов
    for pat in [r"№\s*([A-Za-zА-Яа-я0-9\-_/]{3,})", r"\bN\s*([A-Za-zА-Яа-я0-9\-_/]{3,})"]:
        for m in re.finditer(pat, clean):
            docnum_candidates.append((2, m.group(1)))
    for kw in ["счёт", "счет", "акт", "накладная", "упд", "invoice", "contract", "договор"]:
        for m in re.finditer(kw, clean.lower()):
            left = m.end()
            tail = clean[left:left+30]
            m2 = re.search(r"№\s*([A-Za-zА-Яа-я0-9\-_/]{3,})", tail)
            if m2:
                docnum_candidates.append((3, m2.group(1)))
    if docnum_candidates:
        docnum_candidates.sort(key=lambda x: x[0], reverse=True)
        result["doc_number"] = docnum_candidates[0][1]
    # --- Новый быстрый путь: поиск counterparty ---
    # Ищем строки с ключевыми словами и паттернами организаций
    counterparty_patterns = [
        r"(Общество с ограниченной ответственностью [\"'«][^\"'»]+[\"'»])",
        r"(Акционерное общество [\"'«][^\"'»]+[\"'»])",
        r"(Публичное акционерное общество [\"'«][^\"'»]+[\"'»])",
        r"(Индивидуальный предприниматель [А-ЯЁ][а-яё]+ [А-ЯЁ][а-яё]+)",
        r"(ООО\s+[\"'«][^\"'»]+[\"'»])",
        r"(АО\s+[\"'«][^\"'»]+[\"'»])",
        r"(ПАО\s+[\"'«][^\"'»]+[\"'»])",
        r"(ИП\s+[А-ЯЁ][а-яё]+\s[А-ЯЁ][а-яё]+)",
        r"(Заказчик:?\s*.+)",
        r"(Исполнитель:?\s*.+)",
        r"(Поставщик:?\s*.+)",
        r"(Получатель:?\s*.+)",
        r"(Продавец:?\s*.+)",
        r"(Покупатель:?\s*.+)",
        r"(Контрагент:?\s*.+)"
    ]
    for pat in counterparty_patterns:
        m = re.search(pat, clean)
        if m:
            val = m.group(1).strip()
            # Обрезаем лишнее после запятой/скобки/конца строки
            val = re.split(r'[\n\r,\(\)]', val)[0].strip()
            if 3 < len(val) < 100:
                result["counterparty"] = val
                break
    # --- Новый быстрый путь: поиск contract_number ---
    contract_patterns = [
        r"Договор\s*№\s*([A-Za-zА-Яа-я0-9\-_/]+)",
        r"Contract\s*No\.?\s*([A-Za-zА-Яа-я0-9\-_/]+)",
        r"Соглашение\s*№\s*([A-Za-zА-Яа-я0-9\-_/]+)",
        r"№\s*([A-Za-zА-Яа-я0-9\-_/]+)",
        r"N\s*([A-Za-zА-Яа-я0-9\-_/]+)"
    ]
    for pat in contract_patterns:
        m = re.search(pat, clean)
        if m:
            val = m.group(1).strip()
            if 2 < len(val) < 50:
                result["contract_number"] = val
                break
    # Если хотя бы 3 поля найдены — считаем быстрый путь успешным
    found_fields = sum(1 for v in [result["inn"], result["date"], result["amount"], result["doc_number"], result["counterparty"], result["contract_number"]] if v != "-")
    suspicious = any(is_suspicious(result[f], f) for f in ["amount", "doc_number", "date"])
    if found_fields >= 3 and not suspicious:
        return result

    # --- Медленный путь: LLM ---
    # (остальной код как раньше)
    # Сначала определяем роль нашей компании
    our_role = determine_company_role(clean[:MAX_CHARS])
    # Формируем RAG-контекст
    rag_block = ""
    if rag_context:
        rag_block += "Вот примеры похожих документов:\n"
        for i, frag in enumerate(rag_context, 1):
            rag_block += f"Пример {i}:\n{frag}\n\n"
        rag_block += "----\n"
    # Формируем список нужных полей
    fields_needed = get_fields_for_doc_type(doc_type)
    fields_list = "\n- " + "\n- ".join(fields_needed)
    # Формируем prompt
    prompt = f"""
{rag_block}Это документ типа: {doc_type or '-'}.
Извлеки только следующие поля:{fields_list}
Верни результат в формате JSON с ключами: {', '.join(fields_needed)}.
Текст документа:
"""
    windows = 0
    i = 0
    while i < total_len and windows < 10:
        window_text = clean[i:i+MAX_CHARS]
        full_prompt = prompt + f'"{window_text}"'
        logging.info(f"Prompt to LLM (window {windows+1}, chars {i}-{i+MAX_CHARS}): {full_prompt[:200]}...")
        try:
            response = query_ollama(full_prompt)
            start = response.find('{')
            end = response.rfind('}') + 1
            if start == -1 or end == -1:
                logging.warning("LLM did not return JSON.")
                continue
            json_str = response[start:end]
            fields = json.loads(json_str)
            result = merge_fields(result, fields)
            if all(result.get(k) and result[k] != "-" and result[k].lower() not in ("not specified", "none", "-") for k in fields_needed):
                break
        except Exception as e:
            logging.error(f"Error querying Ollama LLM or parsing JSON: {e}")
            continue
        windows += 1
        i += OVERLAP
    logging.info(f"LLM windows used: {windows}, result: {result}")
    return result

# Пример использования:
# fields = extract_fields_from_text("Текст документа ...")
# print(fields) 